from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "2026-07-23-beyond-rv-sale-terms-draft.md"
OUTPUT = ROOT / "2026-07-23-beyond-rv-sale-terms-draft.docx"

ORANGE = RGBColor(232, 84, 10)
DARK = RGBColor(24, 24, 24)
MUTED = RGBColor(92, 92, 92)
LIGHT_ORANGE = "FFF3EB"
LIGHT_GREY = "F3F4F6"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


def apply_inline(paragraph, text, size=9.2, color=DARK):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, color=color)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.right_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after, color in (
        ("Heading 1", 15, 12, 6, ORANGE),
        ("Heading 2", 11.5, 10, 4, DARK),
        ("Heading 3", 10, 7, 3, DARK),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(9.2)
        style.paragraph_format.left_indent = Mm(7)
        style.paragraph_format.first_line_indent = Mm(-4)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.line_spacing = 1.05

    if "Legal Notice" not in styles:
        notice = styles.add_style("Legal Notice", WD_STYLE_TYPE.PARAGRAPH)
    else:
        notice = styles["Legal Notice"]
    notice.font.name = "Arial"
    notice._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    notice._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    notice.font.size = Pt(9)
    notice.font.bold = True
    notice.font.color.rgb = RGBColor(130, 45, 0)
    notice.paragraph_format.space_after = Pt(0)
    notice.paragraph_format.line_spacing = 1.08

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r1 = hp.add_run("BEYOND RV CAMPERS")
    set_run_font(r1, size=8.5, bold=True, color=ORANGE)
    r2 = hp.add_run("  |  TERMS AND CONDITIONS OF SALE")
    set_run_font(r2, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("LEGAL-REVIEW DRAFT — NOT FOR CUSTOMER USE  •  VERSION 0.1  •  PAGE ")
    set_run_font(fr, size=7.5, bold=True, color=MUTED)
    add_field(fp, "PAGE")
    fr2 = fp.add_run(" OF ")
    set_run_font(fr2, size=7.5, bold=True, color=MUTED)
    add_field(fp, "NUMPAGES")


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("BEYOND RV CAMPERS")
    set_run_font(r, size=11, bold=True, color=ORANGE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    tr = title.add_run("Terms and Conditions of Sale")
    set_run_font(tr, size=24, bold=True, color=DARK)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    mr = meta.add_run("Companion to the Beyond RV Sale Agreement  |  Version 0.1  |  23 July 2026")
    set_run_font(mr, size=9, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Mm(174)
    cell = table.cell(0, 0)
    cell.width = Mm(174)
    set_cell_fill(cell, LIGHT_ORANGE)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p = cell.paragraphs[0]
    p.style = doc.styles["Legal Notice"]
    p.add_run(
        "DO NOT ISSUE TO CUSTOMERS UNTIL AN AUSTRALIAN SOLICITOR HAS REVIEWED AND APPROVED THIS DOCUMENT. "
        "This draft was prepared from the existing Beyond RV 12C Sale Agreement, published Australian RV-sector "
        "terms and current public regulatory guidance. It is not legal advice."
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_markdown(doc, lines):
    i = 0
    ordered_counter = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            i += 1
            continue

        if stripped.startswith("**Legal-review draft"):
            i += 1
            continue

        if stripped.startswith("> **Do not issue"):
            i += 1
            continue

        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                q = lines[i].strip()[1:].strip()
                if q:
                    quote_lines.append(q)
                i += 1
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Mm(174)
            cell = table.cell(0, 0)
            cell.width = Mm(174)
            set_cell_fill(cell, LIGHT_GREY)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            apply_inline(p, " ".join(quote_lines), size=9.2)
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 1")
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:], style="Heading 2")
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            apply_inline(p, stripped[2:])
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number_match and not re.match(r"^\d+\.\d+", stripped):
            p = doc.add_paragraph(style="List Number")
            apply_inline(p, number_match.group(2))
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("> ")
                or nxt.startswith("- ")
                or re.match(r"^(\d+)\.\s+", nxt)
            ):
                break
            paragraph_lines.append(nxt)
            i += 1

        text = " ".join(paragraph_lines)
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        apply_inline(p, text)


def main():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    parse_markdown(doc, lines)

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    core = doc.core_properties
    core.title = "Beyond RV Campers — Terms and Conditions of Sale"
    core.subject = "Legal-review draft companion terms for Beyond RV Sale Agreements"
    core.author = "Passion Industries Pty Ltd trading as Beyond RV Campers"
    core.comments = "Version 0.1 — solicitor review required before customer use"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
